"""Text extraction for uploaded PDF/DOCX files.

Digital PDFs use pdfplumber; DOCX uses python-docx. Scanned / low-text-density
PDFs fall back to page-image rendering + Gemini multimodal OCR.

Pure service — no FastAPI routes, no DB writes. Callers pass bytes + mime_type
(or use extract_text_from_storage to load via storage.read_file).
"""

from __future__ import annotations

import io
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

MIME_PDF = "application/pdf"
MIME_DOCX = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)

# Below this average character count per page, treat the PDF as scanned/sparse.
MIN_CHARS_PER_PAGE = 50
# Cap pages sent to Gemini to keep latency/cost bounded for the hackathon demo.
MAX_OCR_PAGES = 20
GEMINI_OCR_MODEL = "gemini-1.5-flash"


class ExtractionError(Exception):
    """Raised when text cannot be extracted from the given file."""


def extract_text(file_bytes: bytes, mime_type: str) -> str:
    """Extract plain text from PDF or DOCX bytes.

    Digital PDF/DOCX work without GEMINI_API_KEY. Scanned/low-text PDFs require
    settings.gemini_api_key for the multimodal fallback.
    """
    if not file_bytes:
        raise ExtractionError("file_bytes is empty")

    mime = (mime_type or "").strip().lower()
    if mime in (MIME_PDF, "application/x-pdf"):
        return _extract_pdf(file_bytes)
    if mime == MIME_DOCX:
        return _extract_docx(file_bytes)
    raise ExtractionError(
        f"Unsupported mime_type for extraction: {mime_type!r}. "
        "Expected application/pdf or DOCX mime type."
    )


def extract_text_from_storage(storage_key: str, mime_type: str) -> str:
    """Load bytes via storage.read_file, then extract_text."""
    from services import storage

    return extract_text(storage.read_file(storage_key), mime_type)


def _extract_docx(file_bytes: bytes) -> str:
    from docx import Document

    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as exc:
        raise ExtractionError("Failed to open DOCX file") from exc

    parts: list[str] = []
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if text:
            parts.append(text)

    # Tables often hold clause-like content in legal docs.
    for table in doc.tables:
        for row in table.rows:
            cells = [(cell.text or "").strip() for cell in row.cells]
            cells = [c for c in cells if c]
            if cells:
                parts.append(" | ".join(cells))

    text = "\n".join(parts).strip()
    if not text:
        raise ExtractionError("DOCX contained no extractable text")
    return text


def _extract_pdf(file_bytes: bytes) -> str:
    text, page_count = _pdfplumber_text(file_bytes)
    if not _is_low_text_density(text, page_count):
        return text.strip()

    logger.info(
        "PDF text density low (%s chars across %s pages); using Gemini multimodal fallback",
        len(text.strip()),
        page_count,
    )
    images = _render_pdf_page_images(file_bytes)
    return _gemini_ocr_from_images(images)


def _pdfplumber_text(file_bytes: bytes) -> tuple[str, int]:
    import pdfplumber

    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            page_count = len(pdf.pages)
            chunks: list[str] = []
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                if page_text.strip():
                    chunks.append(page_text)
            return "\n\n".join(chunks), page_count
    except ExtractionError:
        raise
    except Exception as exc:
        raise ExtractionError("Failed to open or read PDF with pdfplumber") from exc


def _is_low_text_density(text: str, page_count: int) -> bool:
    """Heuristic: empty or very sparse text suggests a scanned PDF."""
    stripped = (text or "").strip()
    if page_count <= 0:
        return True
    if not stripped:
        return True
    return (len(stripped) / page_count) < MIN_CHARS_PER_PAGE


def _render_pdf_page_images(file_bytes: bytes) -> list[bytes]:
    """Rasterize PDF pages to PNG bytes (requires pypdfium2 via pdfplumber)."""
    import pdfplumber

    images: list[bytes] = []
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            if not pdf.pages:
                raise ExtractionError("PDF has no pages to render for OCR fallback")
            for i, page in enumerate(pdf.pages):
                if i >= MAX_OCR_PAGES:
                    logger.info(
                        "OCR page cap reached (%s); remaining pages skipped",
                        MAX_OCR_PAGES,
                    )
                    break
                try:
                    rendered = page.to_image(resolution=150)
                except Exception as exc:
                    raise ExtractionError(
                        "Failed to render PDF page images for OCR fallback. "
                        "Ensure pypdfium2 is installed (pip install pypdfium2)."
                    ) from exc
                buf = io.BytesIO()
                rendered.original.save(buf, format="PNG")
                images.append(buf.getvalue())
    except ExtractionError:
        raise
    except Exception as exc:
        raise ExtractionError("Failed to render PDF pages for OCR fallback") from exc

    if not images:
        raise ExtractionError("No page images produced for OCR fallback")
    return images


def _require_gemini_api_key() -> str:
    # Lazy import so digital PDF/DOCX paths don't need Gemini configured at call time
    # (settings still loads from .env when config is first imported).
    from config import settings

    key = (settings.gemini_api_key or "").strip()
    if not key:
        raise ExtractionError(
            "GEMINI_API_KEY is required for scanned/low-text PDF extraction fallback. "
            "Digital PDF/DOCX extraction does not need it. "
            "Set GEMINI_API_KEY in lexo/backend/.env."
        )
    return key


def _gemini_ocr_from_images(images: list[bytes]) -> str:
    """Send page PNGs to Gemini multimodal; return extracted plain text."""
    import google.generativeai as genai

    api_key = _require_gemini_api_key()
    # Never log the key or raw page bytes.
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel(GEMINI_OCR_MODEL)

    parts: list[object] = [
        (
            "Extract all readable text from these scanned document page images. "
            "Preserve reading order across pages. Return plain text only — "
            "no commentary, markdown fences, or preamble."
        )
    ]
    for png in images:
        parts.append({"mime_type": "image/png", "data": png})

    try:
        response = model.generate_content(parts)
    except ExtractionError:
        raise
    except Exception as exc:
        logger.exception("Gemini multimodal OCR call failed")
        raise ExtractionError(
            "Gemini multimodal OCR failed. Check GEMINI_API_KEY and model access."
        ) from exc

    text = (getattr(response, "text", None) or "").strip()
    if not text:
        raise ExtractionError("Gemini multimodal OCR returned empty text")
    return text


# ---------------------------------------------------------------------------
# Manual smoke / __main__
# ---------------------------------------------------------------------------

def _minimal_text_pdf(
    message: str = (
        "Lease Agreement Sample Text for Lexo. "
        "This digital PDF contains enough extractable characters "
        "to pass the text-density heuristic without OCR."
    ),
) -> bytes:
    """Tiny digital PDF with embedded text (no reportlab)."""
    # Escape PDF string literals (parentheses / backslash).
    safe = message.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    # Simple one-page PDF; Helvetica + one text show operator.
    stream = f"BT /F1 12 Tf 72 720 Td ({safe}) Tj ET"
    stream_bytes = stream.encode("latin-1")
    objects = [
        b"1 0 obj<< /Type /Catalog /Pages 2 0 R >>endobj\n",
        b"2 0 obj<< /Type /Pages /Kids [3 0 R] /Count 1 >>endobj\n",
        (
            b"3 0 obj<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Contents 4 0 R /Resources<< /Font<< /F1 5 0 R >> >> >>endobj\n"
        ),
        (
            f"4 0 obj<< /Length {len(stream_bytes)} >>stream\n".encode("latin-1")
            + stream_bytes
            + b"\nendstream\nendobj\n"
        ),
        b"5 0 obj<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>endobj\n",
    ]
    out = bytearray(b"%PDF-1.1\n")
    offsets = [0]
    for obj in objects:
        offsets.append(len(out))
        out.extend(obj)
    xref_pos = len(out)
    out.extend(f"xref\n0 {len(offsets)}\n".encode("latin-1"))
    out.extend(b"0000000000 65535 f \n")
    for off in offsets[1:]:
        out.extend(f"{off:010d} 00000 n \n".encode("latin-1"))
    out.extend(
        f"trailer<< /Size {len(offsets)} /Root 1 0 R >>\n"
        f"startxref\n{xref_pos}\n%%EOF\n".encode("latin-1")
    )
    return bytes(out)


def _minimal_empty_pdf() -> bytes:
    """One-page PDF with no text operators — triggers low-density heuristic."""
    objects = [
        b"1 0 obj<< /Type /Catalog /Pages 2 0 R >>endobj\n",
        b"2 0 obj<< /Type /Pages /Kids [3 0 R] /Count 1 >>endobj\n",
        (
            b"3 0 obj<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Contents 4 0 R >>endobj\n"
        ),
        b"4 0 obj<< /Length 0 >>stream\nendstream\nendobj\n",
    ]
    out = bytearray(b"%PDF-1.1\n")
    offsets = [0]
    for obj in objects:
        offsets.append(len(out))
        out.extend(obj)
    xref_pos = len(out)
    out.extend(f"xref\n0 {len(offsets)}\n".encode("latin-1"))
    out.extend(b"0000000000 65535 f \n")
    for off in offsets[1:]:
        out.extend(f"{off:010d} 00000 n \n".encode("latin-1"))
    out.extend(
        f"trailer<< /Size {len(offsets)} /Root 1 0 R >>\n"
        f"startxref\n{xref_pos}\n%%EOF\n".encode("latin-1")
    )
    return bytes(out)


def _minimal_docx(message: str = "Employment Agreement Sample Text for Lexo") -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading("Sample Agreement", level=1)
    doc.add_paragraph(message)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _resolve_sample(path_hint: str | None, default_name: str) -> Path | None:
    """Prefer an uploads/ sample if present; otherwise None (caller uses synthetic)."""
    backend_root = Path(__file__).resolve().parent.parent
    candidates: list[Path] = []
    if path_hint:
        candidates.append(Path(path_hint))
    candidates.extend(
        [
            backend_root / "uploads" / default_name,
            backend_root / default_name,
        ]
    )
    for path in candidates:
        if path.is_file():
            return path
    return None


def _smoke() -> None:
    """Manual smoke: digital PDF/DOCX + mocked scanned fallback.

    Run from lexo/backend:
        python -m services.extraction
        python -m services.extraction path/to/file.pdf
    """
    import sys
    from unittest.mock import patch

    logging.basicConfig(level=logging.INFO)

    arg = sys.argv[1] if len(sys.argv) > 1 else None
    if arg and Path(arg).is_file():
        path = Path(arg)
        data = path.read_bytes()
        mime = MIME_PDF if path.suffix.lower() == ".pdf" else MIME_DOCX
        text = extract_text(data, mime)
        print(f"OK {path.name}: {len(text)} chars")
        print(text[:500])
        return

    pdf_path = _resolve_sample(arg, "sample.pdf")
    docx_path = _resolve_sample(None, "sample.docx")

    pdf_bytes = pdf_path.read_bytes() if pdf_path else _minimal_text_pdf()
    docx_bytes = docx_path.read_bytes() if docx_path else _minimal_docx()
    pdf_label = str(pdf_path) if pdf_path else "<synthetic digital pdf>"
    docx_label = str(docx_path) if docx_path else "<synthetic docx>"

    pdf_text = extract_text(pdf_bytes, MIME_PDF)
    print(f"OK digital PDF ({pdf_label}): {len(pdf_text)} chars")
    print(pdf_text[:200])

    docx_text = extract_text(docx_bytes, MIME_DOCX)
    print(f"OK DOCX ({docx_label}): {len(docx_text)} chars")
    print(docx_text[:200])

    # Simulated scanned PDF: empty page text → must call Gemini path.
    # Mock the Gemini call so this works without a real scanned sample / API hit.
    # Patch via __name__ so `python -m services.extraction` (module=__main__) still works.
    with patch(
        f"{__name__}._gemini_ocr_from_images",
        return_value="OCR_FALLBACK_HIT: scanned page text",
    ) as mock_ocr:
        scanned_text = extract_text(_minimal_empty_pdf(), MIME_PDF)
        assert mock_ocr.called, "expected Gemini multimodal fallback for low-text PDF"
        assert "OCR_FALLBACK_HIT" in scanned_text
    print("OK low-text PDF -> Gemini multimodal fallback (mocked)")


if __name__ == "__main__":
    _smoke()
